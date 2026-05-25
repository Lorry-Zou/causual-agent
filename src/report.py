"""Report generation: Chinese-language Word (.docx) report with charts, tables, and management recommendations."""
from __future__ import annotations

from pathlib import Path
from datetime import datetime
from typing import Any

import numpy as np
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from src.utils import LLMClient, logger, RECOMMENDATIONS_SYSTEM, RECOMMENDATIONS_PROMPT
from src.causal_models import METHODS

# ---------------------------------------------------------------------------
# Style helpers
# ---------------------------------------------------------------------------

_CHART_DIR = Path("output/graphs")


def _set_cn_font(run, font_name: str = "Microsoft YaHei", size: Pt = None, bold: bool = False,
                 color: RGBColor = None) -> None:
    """Set Chinese-compatible font on a run."""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size:
        run.font.size = size
    if bold:
        run.font.bold = bold
    if color:
        run.font.color.rgb = color


def _setup_styles(doc: Document) -> None:
    """Configure document styles for Chinese-language output."""
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5

    for i in range(1, 4):
        hs = doc.styles[f"Heading {i}"]
        hs.font.name = "Microsoft YaHei"
        hs.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        sizes = {1: Pt(18), 2: Pt(15), 3: Pt(13)}
        colors = {1: RGBColor(0x1A, 0x52, 0x76), 2: RGBColor(0x24, 0x71, 0xA3), 3: RGBColor(0x2E, 0x86, 0xC1)}
        hs.font.size = sizes.get(i, Pt(12))
        hs.font.color.rgb = colors.get(i, RGBColor(0x2E, 0x86, 0xC1))
        hs.font.bold = True
        hs.paragraph_format.space_before = Pt(12 if i > 1 else 20)
        hs.paragraph_format.space_after = Pt(6)


def _add_styled_paragraph(doc: Document, text: str, style: str = "Normal",
                          alignment: int = None, bold: bool = False, font_size: Pt = None) -> None:
    """Add a paragraph with Chinese font configured."""
    p = doc.add_paragraph(style=style)
    if alignment is not None:
        p.alignment = alignment
    run = p.add_run(text)
    _set_cn_font(run, size=font_size, bold=bold)
    return p


def _add_table(doc: Document, headers: list[str], rows: list[list[str]],
               col_widths: list[float] = None) -> None:
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        _set_cn_font(run, bold=True, size=Pt(10))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _shade_cell(cell, "1A5276")

    # Data rows
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            _set_cn_font(run, size=Pt(10))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()  # spacing


def _shade_cell(cell, color: str) -> None:
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _add_chart(doc: Document, image_path: str, caption: str = "",
               width: float = 5.5) -> None:
    """Embed a chart image with caption."""
    full_path = Path(image_path)
    if not full_path.is_absolute() and not full_path.exists():
        alt_path = Path("output") / image_path
        if alt_path.exists():
            full_path = alt_path
    if not full_path.exists():
        # Try as relative to cwd
        if not full_path.is_absolute():
            full_path = Path.cwd() / full_path
        if not full_path.exists():
            return
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_img = p_img.add_run()
    run_img.add_picture(str(full_path), width=Inches(width))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap.add_run(caption)
        _set_cn_font(cap_run, size=Pt(9), color=RGBColor(0x7F, 0x8C, 0x8D))
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Section builders
# ---------------------------------------------------------------------------

def _build_section1(doc: Document, intent: dict[str, Any]) -> None:
    """Section 1: User intent analysis."""
    doc.add_heading("1. 用户意图分析", level=2)

    goal = intent.get("business_goal", intent.get("raw_prompt", ""))
    effect = intent.get("effect", intent.get("effect_variable", "未指定"))
    treatment = intent.get("treatment", intent.get("treatment_variable", "未指定"))
    causes = intent.get("candidate_causes", intent.get("causes", []))
    if isinstance(causes, list):
        causes = ", ".join(causes) if causes else "待识别"
    clarification = intent.get("needs_clarification", False)
    questions = intent.get("clarification_questions", "")

    _add_table(doc,
        ["项目", "内容"],
        [
            ["业务目标", goal],
            ["效果变量 (果/因变量)", str(effect)],
            ["处理变量 (因/干预)", str(treatment)],
            ["候选原因", str(causes)],
        ],
        col_widths=[2.0, 4.5],
    )

    if clarification and questions:
        _add_styled_paragraph(doc, f"注意: {questions}", bold=True)


def _build_section2(doc: Document, data_info: dict[str, Any]) -> None:
    """Section 2: Data profile."""
    doc.add_heading("2. 数据概况", level=2)

    sample_size = data_info.get("sample_size", "?")
    feature_count = data_info.get("feature_count", "?")
    data_type = data_info.get("data_type", "?")
    is_panel = "是" if data_info.get("is_panel") else "否"
    n_continuous = data_info.get("n_continuous", "?")
    n_binary = data_info.get("n_binary", "?")
    n_categorical = data_info.get("n_categorical", "?")

    _add_table(doc,
        ["指标", "值"],
        [
            ["样本量", str(sample_size)],
            ["变量总数", str(feature_count)],
            ["连续变量数", str(n_continuous)],
            ["二值变量数", str(n_binary)],
            ["分类变量数", str(n_categorical)],
            ["数据类型", str(data_type)],
            ["是否面板数据", is_panel],
        ],
        col_widths=[2.5, 4.0],
    )

    _add_styled_paragraph(doc, "数据特征描述: " + str(data_info.get("description", "")), font_size=Pt(10))

    # Embed EDA charts
    _add_chart(doc, "output/graphs/correlation_heatmap.png", "图1: 变量相关性热力图")
    _add_chart(doc, "output/graphs/missing_heatmap.png", "图2: 缺失值热力图")


def _build_section3(doc: Document, selected: list[str], method_scores: dict[str, float]) -> None:
    """Section 3: Model selection."""
    doc.add_heading("3. 模型选择", level=2)

    rows = []
    for m in selected:
        info = METHODS.get(m)
        name = info.display_name if info else m
        score = method_scores.get(m, 0)
        prereqs = ", ".join(info.prerequisites) if info and info.prerequisites else "—"
        rows.append([name, f"{score:.1f}", prereqs])

    _add_table(doc,
        ["方法", "得分", "前提条件"],
        rows,
        col_widths=[2.5, 1.0, 3.0],
    )

    _add_styled_paragraph(doc,
        "以上方法基于数据特征（样本量、数据结构、变量类型）自动评分选出。"
        "得分越高表示该方法越适合当前数据。", font_size=Pt(10))


def _build_section4(doc: Document, results: dict[str, Any], output_dir: str) -> None:
    """Section 4: Causal inference results — per-method subsections with charts."""
    doc.add_heading("4. 因果推断结果", level=2)

    chart_map = {
        "panel_regression": "output/graphs/panel_results.png",
        "propensity_score": "output/graphs/psm_distribution.png",
        "diff_in_diff": "output/graphs/did_results.png",
        "instrumental_variable": None,
        "double_ml": "output/graphs/dml_hte.png",
        "rdd": "output/graphs/rdd_plot.png",
        "dag_discovery": "output/graphs/dag.png",
    }

    chart_captions = {
        "panel_regression": "面板回归结果图",
        "propensity_score": "图: 倾向得分匹配前后分布对比",
        "diff_in_diff": "图: 双重差分估计结果",
        "instrumental_variable": "",
        "double_ml": "图: 双机器学习 — 处理效应异质性分布与特征重要性",
        "rdd": "图: 断点回归结果",
        "dag_discovery": "图: 因果有向无环图 (DAG)",
    }

    if not results:
        _add_styled_paragraph(doc, "（未获得因果推断结果）", font_size=Pt(10))
        return

    for method_name, result in results.items():
        info = METHODS.get(method_name)
        display = info.display_name if info else method_name

        doc.add_heading(f"4.{list(results.keys()).index(method_name) + 1} {display}", level=3)

        if result.get("error"):
            _add_styled_paragraph(doc, f"运行失败: {result['error']}", bold=True)
            continue

        ate = result.get("ate", {})
        if ate and ate.get("ate") is not None:
            sig_text = "显著 ***" if ate.get("significant") else "不显著"
            ci_lower = ate.get("ci_lower", float("nan"))
            ci_upper = ate.get("ci_upper", float("nan"))
            _add_table(doc,
                ["指标", "值"],
                [
                    ["处理效应 (ATE)", f"{ate['ate']:.6f}"],
                    ["95% 置信区间", f"[{ci_lower:.4f}, {ci_upper:.4f}]"],
                    ["p值", f"{ate.get('p_value', 'N/A'):.4f}"],
                    ["显著性", sig_text],
                    ["方法", str(ate.get("method", method_name))],
                ],
                col_widths=[2.5, 4.0],
            )

            # Interpretation
            effect_size = ate["ate"]
            abs_eff = abs(effect_size)
            if abs_eff > 0.5:
                magnitude = "大"
            elif abs_eff > 0.1:
                magnitude = "中等"
            else:
                magnitude = "较小"

            if ate.get("significant"):
                interpretation = (
                    f"结果表明，{display}的估计处理效应为 {effect_size:.4f}，"
                    f"效应量{magnitude}，且在统计上显著（p={ate.get('p_value', 0):.4f}）。"
                    f"这意味着处理变量对结果变量存在显著因果影响。"
                )
            else:
                interpretation = (
                    f"结果表明，{display}的估计处理效应为 {effect_size:.4f}，"
                    f"效应量{magnitude}，但统计上不显著（p={ate.get('p_value', 0):.4f}）。"
                    f"因此，现有数据不足以证明处理变量对结果变量存在因果影响。"
                )
            _add_styled_paragraph(doc, interpretation, font_size=Pt(10))

            # Regression coefficient table (full results)
            coef_table = result.get("coef_table", [])
            if coef_table:
                doc.add_paragraph()
                _add_styled_paragraph(doc, "回归系数详情", bold=True, font_size=Pt(11))
                headers = ["变量", "系数", "标准误", "t值", "p值", "95% CI下限", "95% CI上限"]
                rows = []
                for c in coef_table:
                    sig_mark = ""
                    pv = c.get("p_value", 1.0)
                    if not np.isnan(pv):
                        if pv < 0.01:
                            sig_mark = " ***"
                        elif pv < 0.05:
                            sig_mark = " **"
                        elif pv < 0.1:
                            sig_mark = " *"
                    rows.append([
                        c.get("name", ""),
                        f"{c.get('coef', 0):.4f}{sig_mark}",
                        f"{c.get('se', 0):.4f}",
                        f"{c.get('t_stat', 0):.3f}",
                        f"{c.get('p_value', 1.0):.4f}",
                        f"{c.get('ci_lower', 0):.4f}",
                        f"{c.get('ci_upper', 0):.4f}",
                    ])
                _add_table(doc, headers, rows, col_widths=[1.4, 0.9, 0.8, 0.7, 0.7, 0.9, 0.9])

                # Per-variable interpretation
                _add_styled_paragraph(doc, "各变量解读:", bold=True, font_size=Pt(10))
                for c in coef_table:
                    name = c.get("name", "")
                    coef = c.get("coef", 0)
                    pv = c.get("p_value", 1.0)
                    se = c.get("se", 0)
                    if name.startswith("FS: "):  # first-stage instrument, skip detailed interpretation
                        continue
                    direction = "正向" if coef > 0 else "负向"
                    abs_coef = abs(coef)
                    if abs_coef > 0.5:
                        magnitude = "较大"
                    elif abs_coef > 0.1:
                        magnitude = "中等"
                    else:
                        magnitude = "较小"

                    if not np.isnan(pv) and pv < 0.01:
                        sig_desc = "在1%水平上显著"
                    elif not np.isnan(pv) and pv < 0.05:
                        sig_desc = "在5%水平上显著"
                    elif not np.isnan(pv) and pv < 0.1:
                        sig_desc = "在10%水平上边际显著"
                    else:
                        sig_desc = "统计上不显著"

                    interp = (
                        f"• {name}: 系数为 {coef:.4f} (标准误={se:.4f})，{direction}{magnitude}影响，{sig_desc}。"
                    )
                    _add_styled_paragraph(doc, interp, font_size=Pt(9))

        else:
            note = ate.get("note", "") if ate else ""
            n_edges = ate.get("n_edges", 0) if ate else 0
            if n_edges:
                _add_styled_paragraph(doc,
                    f"发现 {n_edges} 条因果边。DAG发现用于识别因果结构而非估计处理效应，"
                    f"建议基于发现的结构选择其他方法进一步估计。", font_size=Pt(10))
            elif note:
                _add_styled_paragraph(doc, str(note), font_size=Pt(10))

        # Diagnostics
        diag = result.get("diagnostics", {})
        if diag:
            diag_rows = []
            for k, v in diag.items():
                if isinstance(v, float):
                    diag_rows.append([k, f"{v:.4f}"])
                elif isinstance(v, bool):
                    diag_rows.append([k, "是" if v else "否"])
                else:
                    diag_rows.append([k, str(v)])
            if diag_rows:
                _add_styled_paragraph(doc, "诊断信息:", bold=True, font_size=Pt(10))
                _add_table(doc, ["诊断项", "值"], diag_rows, col_widths=[2.5, 4.0])

        # Embed chart
        chart_path = chart_map.get(method_name)
        caption = chart_captions.get(method_name, "")
        if chart_path:
            _add_chart(doc, chart_path, caption)

        # Also embed any figure paths returned by the method
        figures = result.get("figures", [])
        if isinstance(figures, list):
            for i, fig_path in enumerate(figures):
                _add_chart(doc, fig_path, f"附图 {i + 1}: {Path(fig_path).stem}")


def _build_section5(doc: Document, llm: LLMClient, intent: dict[str, Any],
                    results: dict[str, Any]) -> None:
    """Section 5: Management recommendations (LLM-generated)."""
    doc.add_heading("5. 管理建议", level=2)

    # Summarize findings for LLM
    findings: list[str] = []
    for method_name, result in results.items():
        info = METHODS.get(method_name)
        display = info.display_name if info else method_name
        ate = result.get("ate", {})
        if ate and ate.get("ate") is not None:
            findings.append(
                f"[{display}] ATE={ate['ate']:.4f} "
                f"(p={ate.get('p_value', 'N/A'):.4f}, "
                f"{'显著' if ate.get('significant') else '不显著'})"
            )
    findings_text = "\n".join(findings) if findings else "暂无量化结果"

    try:
        msgs = [
            {"role": "system", "content": RECOMMENDATIONS_SYSTEM},
            {"role": "user", "content": RECOMMENDATIONS_PROMPT.format(
                business_goal=intent.get("business_goal", intent.get("raw_prompt", "")),
                findings=findings_text,
            )},
        ]
        recs = llm.chat(msgs)
    except Exception:
        recs = "（LLM生成管理建议失败，请基于上述结果自行判断。）"

    _add_styled_paragraph(doc, recs, font_size=Pt(11))

    # Limitations
    doc.add_heading("5.1 分析局限性", level=3)
    limitations = [
        "本分析基于观测数据，虽然采用了因果推断方法控制混淆变量，但仍可能存在未观测的混淆因素。",
        "结果的外部有效性取决于数据代表的总体范围，推广到其他场景时需谨慎。",
        "如果部分方法未达到统计显著水平，建议收集更多数据或考虑其他变量测量方式。",
        "模型假设（如平行趋势、工具变量外生性等）的满足程度会影响因果结论的可信度。",
    ]
    for lim in limitations:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(lim)
        _set_cn_font(run, size=Pt(10))


def _build_footer(doc: Document) -> None:
    """Add generation timestamp footer."""
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"报告生成时间: {now}  |  由AI因果推断分析系统自动生成")
    _set_cn_font(run, size=Pt(8), color=RGBColor(0x95, 0xA5, 0xA6))


# ---------------------------------------------------------------------------
# Main entry
# ---------------------------------------------------------------------------

def generate_report(
    llm: LLMClient,
    output_dir: str,
    intent: dict[str, Any],
    data_info: dict[str, Any],
    selected: list[str],
    method_scores: dict[str, float],
    results: dict[str, Any],
) -> str:
    """Generate a Word (.docx) report with charts, tables, and Chinese narrative. Returns path."""
    _log = logger()
    doc = Document()

    # Page setup — A4
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    _setup_styles(doc)

    # --- Cover / Title ---
    title = doc.add_heading("因果推断分析报告", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_text = intent.get("business_goal", intent.get("raw_prompt", ""))
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(f"—— {subtitle_text}")
    _set_cn_font(sub_run, size=Pt(13), color=RGBColor(0x7F, 0x8C, 0x8D))

    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta_p.add_run(f"生成时间: {now}")
    _set_cn_font(meta_run, size=Pt(9), color=RGBColor(0x95, 0xA5, 0xA6))

    doc.add_paragraph()  # spacer

    # --- Build sections ---
    _build_section1(doc, intent)
    _build_section2(doc, data_info)
    _build_section3(doc, selected, method_scores)
    _build_section4(doc, results, output_dir)
    _build_section5(doc, llm, intent, results)
    _build_footer(doc)

    # --- Save ---
    reports_dir = Path(output_dir) / "reports"
    reports_dir.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    path = reports_dir / f"causal_report_{timestamp}.docx"
    doc.save(str(path))
    _log.info("Word报告已保存: %s", path)
    return str(path)
